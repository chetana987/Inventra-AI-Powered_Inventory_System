from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ---------- STYLES ----------
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ---------- HELPERS ----------
def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()

# ---------- TITLE PAGE ----------
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Inventra – AI-Powered Inventory Management System')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Complete Project Documentation & Interview Preparation Guide')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Prepared for: Chetana Mahajan\nDate: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)

doc.add_page_break()

# ========================================================================
# SECTION 1: COMPLETE PROJECT DOCUMENTATION
# ========================================================================
doc.add_heading('1. Complete Project Documentation', level=1)

doc.add_heading('1.1 Project Overview', level=2)
doc.add_paragraph(
    'Inventra is a full-stack AI-powered inventory management system built with Spring Boot 3 and Angular 21. '
    'It enables businesses to track stock levels, manage products, record inventory transactions, '
    'and get actionable insights through a natural language AI assistant powered by Google Gemini API. '
    'The application is containerized with Docker and deployed on Railway with PostgreSQL as the primary database '
    'and Redis for caching.'
)

doc.add_heading('1.2 Tech Stack', level=2)
add_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend', 'Angular 21, TypeScript, SCSS, RxJS', 'Single-page application with reactive UI'],
        ['Backend', 'Java 21, Spring Boot 3.3.5', 'REST API and business logic'],
        ['Security', 'Spring Security, JWT (jjwt 0.12.5), BCrypt', 'Authentication and authorization'],
        ['Database', 'PostgreSQL (primary), Redis (caching)', 'Persistent storage and caching'],
        ['AI', 'Google Gemini API (gemini-2.0-flash)', 'Natural language query processing'],
        ['Build', 'Maven, Docker, Docker Compose', 'Build and containerization'],
        ['Deploy', 'Railway', 'Cloud hosting and PostgreSQL plugin'],
        ['Docs', 'SpringDoc OpenAPI (Swagger)', 'API documentation'],
    ]
)

doc.add_heading('1.3 Project Structure', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'inventra/\n'
    '├── frontend/                    # Angular SPA\n'
    '│   ├── src/app/\n'
    '│   │   ├── components/          # Reusable UI (skeleton, empty-state, toast, footer)\n'
    '│   │   ├── guards/              # authGuard – route protection\n'
    '│   │   ├── interceptors/        # jwtInterceptor, timeoutInterceptor\n'
    '│   │   ├── models/              # TypeScript interfaces\n'
    '│   │   ├── pages/               # Login, Signup, Dashboard, Products, etc.\n'
    '│   │   └── services/            # Auth, Product, Inventory, Dashboard, AI, Toast\n'
    '├── src/main/java/com/inventra/\n'
    '│   ├── config/                  # DataSeeder, WebConfig, RedisConfig, OpenApiConfig, AuditAspect\n'
    '│   ├── controller/              # REST controllers (7 controllers)\n'
    '│   ├── dto/                     # Request/Response DTOs\n'
    '│   ├── entity/                  # JPA entities (Product, User, InventoryTransaction, AuditLog)\n'
    '│   ├── exception/               # GlobalExceptionHandler, ApiException, BadRequestException, ResourceNotFoundException\n'
    '│   ├── repository/              # JPA repositories (4 repositories)\n'
    '│   ├── security/                # SecurityConfig, JwtAuthFilter, JwtService, UserDetailsServiceImpl\n'
    '│   └── service/                 # Business logic (8 services)\n'
    '├── Dockerfile                   # Multi-stage build (Node -> Maven -> JRE)\n'
    '├── docker-compose.yml\n'
    '├── pom.xml\n'
    '└── .env.example\n'
)
run.font.size = Pt(9)
run.font.name = 'Consolas'

doc.add_heading('1.4 Key Features', level=2)
features = [
    'Dashboard: Real-time KPIs, donut chart (category distribution), bar chart (monthly trends), recent transactions',
    'Product Management: Full CRUD with pagination, sorting, multi-field search, and filtering by category',
    'Stock Transactions: Stock-in and stock-out with pessimistic locking for data integrity',
    'Low-Stock Alerts: Automatic detection of products below minimum stock levels; daily cron job at midnight',
    'AI Assistant: Natural language queries via Google Gemini API (e.g., "Show me low stock items", "What is the total inventory value?")',
    'Role-Based Access: Admin (full access) and Staff (view+ transactions) roles',
    'JWT Authentication: Token-based auth with 24-hour expiry, stored in localStorage/sessionStorage',
    'Audit Logging: AOP-based logging for all product and inventory operations',
    'Caching: Redis (or simple in-memory fallback) for dashboard, product list, and low-stock queries',
    'Responsive UI: Loading skeletons, empty states, toast notifications, dark sidebar layout',
]
for f in features:
    add_bullet(doc, f)

doc.add_heading('1.5 Module Descriptions', level=2)

modules = [
    ('Auth Module', 'Handles user registration and login. Uses BCrypt for password hashing and JWT for token generation. '
     'The first registered user gets ADMIN role; subsequent users get STAFF. Tokens expire in 24 hours.'),
    ('Product Module', 'Manages product catalog with CRUD operations. Supports paginated listing with sorting (by id, name, price, quantity) '
     'and filtering (by name with LIKE, by category with exact match). Low-stock detection compares quantity against minimumStockLevel.'),
    ('Inventory Module', 'Records stock-in and stock-out transactions. Uses pessimistic write locks to prevent race conditions during concurrent stock updates. '
     'Validates sufficient stock before stock-out. Maintains complete audit trail with timestamps and user attribution.'),
    ('Dashboard Module', 'Aggregates data from multiple sources: total products, total stock quantity, low-stock count, '
     'category distribution (for donut chart), monthly transaction trends (for bar chart), and recent transactions list. '
     'Results are cached with Redis for 10 minutes.'),
    ('AI Assistant Module', 'Accepts natural language questions, sends them to Google Gemini API with inventory context, '
     'classifies the intent (summary, low-stock, stock-value, category-breakdown, etc.), executes the appropriate database query, '
     'and returns a summary with supporting data. Has a local rule-based fallback if Gemini API is unavailable.'),
    ('User Module', 'Profile management: view profile, update name/email, change password. Requires current password verification for password changes.'),
    ('Health Module', 'Simple health check endpoint used by Docker HEALTHCHECK and load balancers. Returns application name, status, and timestamp.'),
    ('Audit Module', 'AOP aspect that intercepts product and inventory service methods. Logs who performed what action with timestamp.'),
]
for name, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ========================================================================
# SECTION 2: DATABASE SCHEMA DIAGRAM
# ========================================================================
doc.add_heading('2. Database Schema Diagram', level=1)

doc.add_heading('2.1 Entity Relationship Diagram (Textual)', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '┌─────────────────────────────────────────────────────────────────────┐\n'
    '│                      INVENTRA DATABASE SCHEMA                       │\n'
    '├─────────────────────────────────────────────────────────────────────┤\n'
    '│                                                                     │\n'
    '│  ┌──────────────────────┐        ┌────────────────────────────┐    │\n'
    '│  │       users          │        │       products              │    │\n'
    '│  ├──────────────────────┤        ├────────────────────────────┤    │\n'
    '│  │ id (PK)         BIGINT│<────┐│ id (PK)               BIGINT│    │\n'
    '│  │ name            VARCHAR│    ││ product_code    VARCHAR(50) │    │\n'
    '│  │ email           VARCHAR│    ││ name             VARCHAR(150)│    │\n'
    '│  │ password        VARCHAR│    ││ category         VARCHAR(100)│    │\n'
    '│  │ role         VARCHAR(20)│   ││ description          TEXT   │    │\n'
    '│  │ created_at   TIMESTAMP │    ││ price         DECIMAL(12,2) │    │\n'
    '│  └──────────────────────┘    ││ quantity           INTEGER   │    │\n'
    '│                              ││ minimum_stock_level INTEGER  │    │\n'
    '│                              ││ created_at       TIMESTAMP   │    │\n'
    '│                              ││ updated_at       TIMESTAMP   │    │\n'
    '│                              │└──────────┬───────────────────┘    │\n'
    '│                              │           │                         │\n'
    '│  ┌──────────────────────┐    │           │                         │\n'
    '│  │  inventory_transactions │  │           │                         │\n'
    '│  ├──────────────────────┤    │           │                         │\n'
    '│  │ id (PK)            BIGINT│  │           │                         │\n'
    '│  │ product_id (FK)    BIGINT│──┘           │                         │\n'
    '│  │ transaction_type VARCHAR │               │                         │\n'
    '│  │ quantity            INTEGER│             │                         │\n'
    '│  │ remarks               TEXT│             │                         │\n'
    '│  │ transaction_date TIMESTAMP│             │                         │\n'
    '│  │ performed_by (FK)  BIGINT│─────────────┘                         │\n'
    '│  └──────────────────────┘                                           │\n'
    '│                                                                     │\n'
    '│  ┌──────────────────────┐                                           │\n'
    '│  │     audit_logs       │                                           │\n'
    '│  ├──────────────────────┤                                           │\n'
    '│  │ id (PK)            BIGINT│                                       │\n'
    '│  │ username      VARCHAR(100)│                                      │\n'
    '│  │ action        VARCHAR(50)│                                       │\n'
    '│  │ timestamp       TIMESTAMP│                                       │\n'
    '│  └──────────────────────┘                                           │\n'
    '└─────────────────────────────────────────────────────────────────────┘\n'
)
run.font.size = Pt(8)
run.font.name = 'Consolas'

doc.add_heading('2.2 Table Details', level=2)

add_table(doc,
    ['Table', 'Columns', 'Constraints', 'Relationships'],
    [
        ['users', 'id, name, email, password, role, created_at', 'PK(id), UQ(email)', 'Referenced by inventory_transactions.performed_by'],
        ['products', 'id, product_code, name, category, description, price, quantity, minimum_stock_level, created_at, updated_at', 'PK(id), UQ(product_code)', 'Referenced by inventory_transactions.product_id'],
        ['inventory_transactions', 'id, product_id, transaction_type, quantity, remarks, transaction_date, performed_by', 'PK(id)', 'FK(product_id -> products.id), FK(performed_by -> users.id)'],
        ['audit_logs', 'id, username, action, timestamp', 'PK(id)', 'None (standalone)'],
    ]
)

doc.add_heading('2.3 Entity Fields Explained', level=2)

entities_to_explain = [
    ('User', [
        ('id', 'Primary key, auto-generated identity'),
        ('name', "User's full name"),
        ('email', 'Login email, must be unique'),
        ('password', 'BCrypt-hashed password'),
        ('role', 'Enum: ADMIN or STAFF'),
        ('createdAt', 'Account creation timestamp'),
    ]),
    ('Product', [
        ('id', 'Primary key, auto-generated identity'),
        ('productCode', 'SKU/product code, must be unique'),
        ('name', 'Product name'),
        ('category', 'Product category (Electronics, Furniture, Stationery)'),
        ('description', 'Detailed description (TEXT type)'),
        ('price', 'Unit price (DECIMAL 12,2)'),
        ('quantity', 'Current stock quantity'),
        ('minimumStockLevel', 'Threshold below which product is considered low-stock'),
        ('createdAt', 'Record creation timestamp'),
        ('updatedAt', 'Last update timestamp, auto-set via @PreUpdate'),
    ]),
    ('InventoryTransaction', [
        ('id', 'Primary key, auto-generated identity'),
        ('product', 'Many-to-one relationship to Product'),
        ('transactionType', 'Enum: STOCK_IN or STOCK_OUT'),
        ('quantity', 'Quantity moved in this transaction'),
        ('remarks', 'Optional note/reason'),
        ('transactionDate', 'When the transaction occurred'),
        ('performedBy', 'Many-to-one relationship to User'),
    ]),
]
for entity, fields in entities_to_explain:
    p = doc.add_paragraph()
    run = p.add_run(f'{entity}:')
    run.bold = True
    add_table(doc, ['Field', 'Description'], fields)

doc.add_page_break()

# ========================================================================
# SECTION 3: API DOCUMENTATION
# ========================================================================
doc.add_heading('3. API Documentation', level=1)

doc.add_heading('3.1 Authentication APIs', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Request Body', 'Response', 'Description'],
    [
        ['POST', '/api/auth/register', 'Public', 'RegisterRequest {name, email, password}', 'ApiResponse<AuthResponse> {token, name, email, role}', 'Register new user (1st = ADMIN, rest = STAFF)'],
        ['POST', '/api/auth/login', 'Public', 'LoginRequest {email, password}', 'ApiResponse<AuthResponse> {token, name, email, role}', 'Login with email/password, returns JWT'],
    ]
)

doc.add_heading('3.2 Product APIs', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Params', 'Description'],
    [
        ['GET', '/api/products', 'Authenticated', 'page, size, sortBy, sortDir, name?, category?', 'Paginated product list with search/filter'],
        ['GET', '/api/products/{id}', 'Authenticated', 'Path: id', 'Get single product by ID'],
        ['GET', '/api/products/low-stock', 'Authenticated', 'page, size', 'Paginated low-stock products'],
        ['POST', '/api/products', 'Admin', 'ProductRequest body', 'Create new product'],
        ['PUT', '/api/products/{id}', 'Admin', 'Path: id, ProductRequest body', 'Update existing product'],
        ['DELETE', '/api/products/{id}', 'Admin', 'Path: id', 'Delete product'],
    ]
)

doc.add_heading('3.3 Inventory APIs', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Params', 'Description'],
    [
        ['POST', '/api/inventory/stock-in', 'Admin', 'InventoryRequest {productId, quantity, remarks}', 'Add stock (pessimistic lock)'],
        ['POST', '/api/inventory/stock-out', 'Admin', 'InventoryRequest {productId, quantity, remarks}', 'Remove stock (validates sufficient qty)'],
        ['GET', '/api/inventory/history', 'Authenticated', 'productId?, page, size', 'Transaction history with optional product filter'],
    ]
)

doc.add_heading('3.4 Dashboard API', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['GET', '/api/dashboard', 'Authenticated', 'Returns totalProducts, totalStockQuantity, lowStockCount, category distribution, monthly trends, recent transactions'],
    ]
)

doc.add_heading('3.5 AI Assistant API', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Request Body', 'Description'],
    [
        ['POST', '/api/ai/query', 'Authenticated', 'AiQueryRequest {question}', 'Natural language query, returns intent + summary + data'],
    ]
)

doc.add_heading('3.6 User Profile APIs', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['GET', '/api/users/me', 'Authenticated', 'Get current user profile'],
        ['PUT', '/api/users/profile', 'Authenticated', 'Update name and email'],
        ['PUT', '/api/users/password', 'Authenticated', 'Change password (requires current password)'],
    ]
)

doc.add_heading('3.7 Health API', level=2)
add_table(doc,
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['GET', '/api/health', 'Public', 'Returns app name, status UP, and timestamp'],
    ]
)

doc.add_heading('3.8 API Response Format', level=2)
doc.add_paragraph('All APIs return responses in a standard wrapper:')
p = doc.add_paragraph()
run = p.add_run(
    '{\n'
    '  "status": 200,\n'
    '  "message": "Success",\n'
    '  "data": { ... },    // Type varies by endpoint\n'
    '  "timestamp": "2026-06-07T04:06:16"  // ISO-8601\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('Error responses follow the same structure:')
p = doc.add_paragraph()
run = p.add_run(
    '{\n'
    '  "status": 400,\n'
    '  "message": "Email already registered",\n'
    '  "data": null,\n'
    '  "timestamp": "2026-06-07T04:06:16"\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_page_break()

# ========================================================================
# SECTION 4: ARCHITECTURE DIAGRAM
# ========================================================================
doc.add_heading('4. Architecture Diagram', level=1)

doc.add_heading('4.1 System Architecture', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '┌────────────────────────────────────────────────────────────────────────────┐\n'
    '│                            CLIENT (Browser)                                │\n'
    '│  ┌──────────┐ ┌────────────┐ ┌──────────────┐ ┌────────────────────────┐ │\n'
    '│  │ Landing  │ │ Auth Pages │ │Dashboard Views│ │  AI Assistant Chat UI │ │\n'
    '│  │ Page     │ │(Login/Sign)│ │(Charts/Tables)│ │  (Chat Interface)     │ │\n'
    '│  └──────────┘ └────────────┘ └──────────────┘ └────────────────────────┘ │\n'
    '└──────────────────────────┬─────────────────────────────────────────────────┘\n'
    '                           │\n'
    '                    HTTP Requests\n'
    '                    + JWT Token (Authorization: Bearer <token>)\n'
    '                           │\n'
    '┌──────────────────────────▼─────────────────────────────────────────────────┐\n'
    '│                     SPRING BOOT BACKEND (Port 8080)                        │\n'
    '│                                                                           │\n'
    '│  ┌──────────────────────────────────────────────────────────────────────┐ │\n'
    '│  │                   SECURITY FILTER CHAIN                              │ │\n'
    '│  │  JwtAuthFilter -> SecurityContext -> AuthorizationFilter             │ │\n'
    '│  └──────────────────────────────────────────────────────────────────────┘ │\n'
    '│                                                                           │\n'
    '│  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────────┐ │\n'
    '│  │ AuthController│ │ProductControl│ │InventoryCont │ │  AiController   │ │\n'
    '│  │ /api/auth    │ │ /api/products│ │ /api/invento.│ │  /api/ai        │ │\n'
    '│  └──────┬───────┘ └──────┬───────┘ └──────┬───────┘ └───────┬──────────┘ │\n'
    '│         │                │                │                 │            │\n'
    '│  ┌──────▼───────┐ ┌──────▼───────┐ ┌──────▼───────┐ ┌──────▼──────────┐ │\n'
    '│  │  AuthService  │ │ProductService │ │InventoryServ │ │  AiService      │ │\n'
    '│  │  + JwtService │ │+Cacheable    │ │+Pessimistic  │ │  + GeminiService │ │\n'
    '│  └──────┬───────┘ └──────┬───────┘ │  Lock        │ └──────┬──────────┘ │\n'
    '│         │                │         └──────┬───────┘        │            │\n'
    '│  ┌──────▼───────┐ ┌──────▼───────┐        │        ┌──────▼──────────┐ │\n'
    '│  │ UserRepo     │ │ProductRepo   │  InventoryTxRepo │ (Google Gemini)│ │\n'
    '│  └──────┬───────┘ └──────┬───────┘        │        └─────────────────┘ │\n'
    '│         │                │                 │                            │\n'
    '│  ┌──────▼────────────────▼────────────────▼──────────────────────────┐ │\n'
    '│  │                    SERVICE LAYER + DAO LAYER                       │ │\n'
    '│  │  AuditAspect (@Aspect) -> AuditLogService -> AuditLogRepository   │ │\n'
    '│  │  LowStockAlertService (@Scheduled cron)                          │ │\n'
    '│  │  DashboardService (@Cacheable)                                   │ │\n'
    '│  └──────────────────────────────────────────────────────────────────┘ │\n'
    '└──────────────────┬───────────────────────────────────────┬──────────────┘\n'
    '                   │                                       │\n'
    '          ┌────────▼────────┐                    ┌──────────▼──────────┐\n'
    '          │   PostgreSQL    │                    │   Redis Cache       │\n'
    '          │  (Primary DB)   │                    │  (Dashboard/Products)│\n'
    '          └─────────────────┘                    └─────────────────────┘\n'
)
run.font.size = Pt(7)
run.font.name = 'Consolas'

doc.add_heading('4.2 Request Lifecycle', level=2)
steps = [
    'Browser makes HTTP request (e.g., GET /api/products)',
    'Request passes through Spring Security filter chain: CORS -> JwtAuthFilter -> SecurityContext -> Authorization',
    'JwtAuthFilter extracts JWT from Authorization header, validates it, creates SecurityContext',
    'AuthorizationFilter checks URL pattern against configured rules (permitAll, authenticated, hasRole)',
    'Request reaches DispatcherServlet which routes to the correct controller method',
    'Controller calls Service layer which contains business logic',
    'Service calls Repository (JPA) to query/update the database',
    'For dashboard/products: Redis cache is checked first (if enabled)',
    'For stock transactions: Pessimistic write lock is acquired on the product row',
    'Response flows back through the filter chain and is returned to the browser',
]
for i, step in enumerate(steps, 1):
    add_bullet(doc, step, f'{i}. ')

doc.add_heading('4.3 Authentication Flow', level=2)
steps = [
    'User enters email/password on Login page',
    'Frontend sends POST /api/auth/login with credentials',
    'AuthService.login() calls AuthenticationManager.authenticate()',
    'DaoAuthenticationProvider uses UserDetailsServiceImpl.loadUserByUsername() to fetch user',
    'BCryptPasswordEncoder verifies the password hash',
    'If valid: JwtService.generateToken() creates a JWT with email, role, 24h expiry',
    'Response returns { token, name, email, role }',
    'Frontend stores token in localStorage/sessionStorage',
    'Subsequent API requests include Authorization: Bearer <token> header via JwtInterceptor',
    'JwtAuthenticationFilter validates token on every request and sets SecurityContext',
]
for i, step in enumerate(steps, 1):
    add_bullet(doc, step, f'{i}. ')

doc.add_page_break()

# ========================================================================
# SECTION 5: 50 INTERVIEW QUESTIONS AND ANSWERS
# ========================================================================
doc.add_heading('5. 50 Interview Questions and Answers', level=1)

questions = [
    # General / Project Overview
    ("Q1: What is Inventra and what problem does it solve?",
     "Inventra is an AI-powered inventory management system. It solves the problem of manual inventory tracking by providing "
     "a centralized platform where businesses can manage products, record stock movements, monitor low-stock alerts, and "
     "get actionable insights through natural language queries powered by AI."),
    
    ("Q2: What is the tech stack used in Inventra?",
     "Frontend: Angular 21 with TypeScript and SCSS. Backend: Java 21 with Spring Boot 3.3.5. Database: PostgreSQL. "
     "Caching: Redis. AI: Google Gemini API. Authentication: JWT with BCrypt. Deployment: Docker on Railway."),
    
    ("Q3: What are the main features of the application?",
     "Dashboard with KPIs and charts, Product CRUD with pagination, Stock-in/Stock-out transactions with audit trail, "
     "Low-stock alerts with daily cron job, AI Assistant for natural language queries, Role-based access (Admin/Staff), "
     "JWT authentication, and Audit logging."),

    # Architecture & Design
    ("Q4: Explain the architecture of the project.",
     "The project follows a client-server architecture with a REST API. The Angular frontend communicates with the Spring Boot backend "
     "via HTTP requests. The backend has a layered architecture: Controller (REST endpoints) -> Service (business logic) -> "
     "Repository (data access). Spring Security with JWT handles authentication. Redis provides caching. The AI module "
     "integrates with Google Gemini API."),
    
    ("Q5: Why did you choose PostgreSQL over MySQL?",
     "PostgreSQL was chosen because Railway (our deployment platform) provides a managed PostgreSQL plugin with automatic "
     "connection string injection via environment variables. PostgreSQL also offers advanced features like JSONB, "
     "superior concurrency, and better support for complex queries used in our reporting/dashboard features."),
    
    ("Q6: Why did you use Redis? What is cached?",
     "Redis is used for caching frequently accessed, slowly changing data. The dashboard aggregate data, paginated product "
     "lists, and low-stock product lists are cached with a 10-minute TTL. This reduces database load and improves response times."),
    
    ("Q7: Explain the JWT authentication flow in detail.",
     "When a user logs in, the server validates credentials and generates a JWT containing the user's email, role, "
     "issue timestamp, and expiration (24 hours). The token is signed with HMAC-SHA256 using a secret key. "
     "The frontend stores this token and sends it as an Authorization: Bearer header on all subsequent requests. "
     "A JwtAuthenticationFilter intercepts each request, validates the token, and sets the SecurityContext."),

    # Spring Boot
    ("Q8: What is the role of @SpringBootApplication?",
     "@SpringBootApplication is a convenience annotation that combines @Configuration, @EnableAutoConfiguration, and "
     "@ComponentScan. It tells Spring Boot to auto-configure the application, scan for components in the current package "
     "and sub-packages, and treat the class as a source of bean definitions."),
    
    ("Q9: What design patterns are used in this project?",
     "Several patterns: (1) Singleton – Spring beans are singletons by default. (2) Factory – Spring creates and manages beans. "
     "(3) Proxy – AOP uses JDK/CGLIB proxies for transactions, caching, and aspect logging. "
     "(4) Builder – Lombok @Builder on entities and DTOs. (5) DTO – Separate request/response objects from entities. "
     "(6) Repository – Spring Data JPA repositories abstract data access."),
    
    ("Q10: Explain the use of @Transactional in this project.",
     "@Transactional is used on service methods to ensure database operations are atomic. For example, when processing "
     "a stock-out transaction, we need to both create a transaction record AND update the product quantity in a single "
     "database transaction. If either fails, both are rolled back. readOnly=true is used on read-only methods for optimization."),
    
    ("Q11: What is the difference between @Component, @Service, and @Repository?",
     "@Component is a generic stereotype. @Service indicates business logic layer. @Repository indicates data access layer "
     "and Spring automatically adds persistence exception translation. All three enable component scanning. "
     "In Inventra, we use @Service for service classes and @Repository (via extending JpaRepository) for data access."),
    
    ("Q12: How does dependency injection work in this project?",
     "We use constructor injection via Lombok's @RequiredArgsConstructor. Spring automatically injects dependencies "
     "through the constructor. For example, ProductController receives ProductService through its constructor. "
     "This makes testing easier (can pass mocks) and ensures immutability."),
    
    ("Q13: What is the purpose of application.yml and profile-specific YAML files?",
     "application.yml contains default configuration. Profile-specific files (application-dev.yml, application-railway.yml, "
     "application-test.yml) override default settings for specific environments. Railway uses the 'railway' profile for "
     "PostgreSQL config, test uses H2 in-memory database, and dev uses local MySQL."),
    
    ("Q14: Explain the exception handling mechanism.",
     "We use @RestControllerAdvice on GlobalExceptionHandler class. It has handler methods for different exception types: "
     "ApiException returns 400/404, MethodArgumentNotValidException returns 400 with field errors, "
     "BadCredentialsException returns 401, AccessDeniedException returns 403, and generic Exception returns 500. "
     "All responses follow the standard ApiResponse format."),
    
    ("Q15: How does the pagination work in the product listing?",
     "The frontend sends page, size, sortBy, sortDir parameters. The controller uses @RequestParam with defaults "
     "(page=0, size=10, sortBy=id, sortDir=asc). The service creates a Pageable object and passes it to the "
     "repository. Spring Data JPA generates the appropriate LIMIT, OFFSET, and ORDER BY SQL clauses. "
     "The response includes content, page number, total pages, and total elements."),

    # JPA & Hibernate
    ("Q16: Why did you use ddl-auto=update? Is it safe for production?",
     "ddl-auto=update tells Hibernate to automatically create or update database tables based on entity definitions. "
     "It's convenient for development but NOT recommended for production because it can make destructive changes. "
     "The production profile uses ddl-auto=validate which only checks that the schema matches entities and fails if not."),
    
    ("Q17: Explain lazy loading in this project.",
     "InventoryTransaction has @ManyToOne(fetch=FetchType.LAZY) relationships to Product and User. Lazy loading means "
     "the related entities are not fetched from the database until they are accessed. This improves performance by "
     "avoiding unnecessary JOINs. However, accessing lazy-loaded entities outside a transaction can cause "
     "LazyInitializationException, so we use JOIN FETCH in queries that need the related data."),
    
    ("Q18: What is pessimistic locking and where is it used?",
     "Pessimistic locking prevents concurrent access to a database row by locking it. In this project, "
     "the stock-in and stock-out operations use @Lock(PESSIMISTIC_WRITE) on the Product entity. "
     "This ensures that when one user is updating stock quantity, another user cannot read or update "
     "the same product until the transaction completes."),
    
    ("Q19: How did you optimize database queries?",
     "Key optimizations: (1) JOIN FETCH in InventoryTransactionRepository to avoid N+1 queries. "
     "(2) Batch inserts with hibernate.jdbc.batch_size=25 and order_inserts/order_updates=true. "
     "(3) Cacheable annotations for dashboard and product lists. (4) Pagination to avoid loading all records. "
     "(5) Specific JPQL queries instead of deriving them from method names."),

    # Security
    ("Q20: How is CSRF protection handled?",
     "CSRF (Cross-Site Request Forgery) protection is disabled because the API uses JWT tokens which are immune to "
     "CSRF attacks. JWT tokens are sent via Authorization headers (not cookies), so malicious sites cannot forge requests."),
    
    ("Q21: How are passwords stored securely?",
     "Passwords are hashed using BCryptPasswordEncoder, which implements bcrypt strong hashing with built-in salt. "
     "BCrypt is intentionally slow (configurable strength/cost factor) to resist brute-force attacks. "
     "Plain-text passwords are never stored or logged."),
    
    ("Q22: Explain CORS configuration and why it's needed.",
     "CORS (Cross-Origin Resource Sharing) controls which origins can access the API. In development, the Angular app "
     "runs on localhost:4200 while the backend runs on localhost:8080, so CORS headers are needed. In production on Railway, "
     "both are served from the same origin. The configuration reads allowed origins from application.yml."),
    
    ("Q23: How did you secure different endpoints for different roles?",
     "Endpoint security is configured in SecurityConfig's SecurityFilterChain using requestMatchers. "
     "/api/admin/** requires ADMIN role. /api/auth/** and /api/health are public. All other /api/** endpoints "
     "require authentication. Additionally, @PreAuthorize("hasRole('ADMIN')") is used on specific controller methods "
     "like product creation and stock transactions."),

    # Angular / Frontend
    ("Q24: How does the JWT interceptor work?",
     "The jwtInterceptor is an Angular HttpInterceptorFn that runs before every HTTP request. It injects AuthService, "
     "calls getToken() to retrieve the JWT from localStorage, and clones the request with an Authorization: Bearer <token> "
     "header if the token exists. This ensures all API calls include the authentication token automatically."),
    
    ("Q25: How is the auth guard implemented?",
     "The authGuard implements CanActivate. It calls AuthService.isLoggedIn(), which checks if a token exists and "
     "if it's not expired (by decoding the JWT payload and checking the exp claim). If valid, it returns true. "
     "Otherwise, it redirects to the /login page using Router.navigate()."),
    
    ("Q26: Explain the timeout interceptor.",
     "The timeoutInterceptor applies an RxJS timeout(15000) operator to all HTTP requests. If a request takes more than "
     "15 seconds, it throws a TimeoutError. The component's error handler catches this and displays a user-friendly "
     "error message. This prevents the UI from hanging indefinitely on slow network connections."),
    
    ("Q27: How is form validation handled on the frontend?",
     "Forms use Angular Template-Driven or Reactive Forms with both HTML5 validation (required, minlength, pattern) "
     "and custom validators (password match for signup). The signup page has a password strength meter. "
     "The backend also validates with Jakarta Validation annotations (@NotBlank, @Email, @Size) as a second layer."),
    
    ("Q28: How does the signup page handle the confirm password field?",
     "The signup form includes a confirmPassword field for UX. Before sending the request, the component "
     "destructures and removes confirmPassword from the payload: const { confirmPassword, ...payload } = this.signupForm.getRawValue(). "
     "Only name, email, and password are sent to the backend. The confirmPassword matching is validated on the client side."),

    # AI Integration
    ("Q29: How does the AI Assistant work?",
     "The user types a natural language question in the chat interface. The frontend sends it to POST /api/ai/query. "
     "The AiService first tries to classify the intent using Google Gemini API. If Gemini is unavailable, "
     "a local rule-based classifier checks keywords. Based on the intent (summary, low-stock, stock-value, "
     "category-breakdown, etc.), the service executes the appropriate database query and returns a summary with data."),
    
    ("Q30: What happens if the Gemini API is down?",
     "The GeminiService has a fallback mechanism. If the Gemini API call fails (timeout or network error), "
     "it catches the exception and calls a local classifyIntentLocally() method that uses simple keyword matching "
     "to determine the intent. This ensures the AI feature degrades gracefully."),

    # Docker & Deployment
    ("Q31: Explain the multi-stage Docker build.",
     "The Dockerfile has 3 stages: (1) Frontend build – uses Node 20 to build Angular with production config. "
     "(2) Backend build – uses Maven 3.9 + Java 21 to compile the Spring Boot app and copies the Angular build "
     "output into src/main/resources/static. (3) Runtime – uses a slim Java 21 JRE image, copies the JAR, "
     "and runs it. This keeps the final image small (only runtime dependencies)."),

    ("Q32: How did you deploy on Railway?",
     "The GitHub repository is connected to Railway. On push, Railway detects the Dockerfile and builds the image. "
     "A PostgreSQL plugin is added which injects PGHOST, PGPORT, PGDATABASE, PGUSER, PGPASSWORD environment variables. "
     "The Railway profile in application-railway.yml reads these env vars to configure the datasource."),
]

# Interview Q&A - rest
questions += [
    # Performance & Optimization
    ("Q33: How did you handle the N+1 query problem?",
     "The N+1 problem occurs when you fetch entities and then lazily load related entities one by one. "
     "In InventoryTransactionRepository, the queries use JOIN FETCH to eagerly load both product and user "
     "relationships in a single query. This eliminates the N+1 problem for transaction history and recent transactions."),
    
    ("Q34: What caching strategy is used?",
     "We use a read-through caching strategy with @Cacheable annotations. The cache is checked before executing "
     "the method. If a cached value exists, it's returned directly. Otherwise, the method executes, and the result "
     "is cached. @CacheEvict invalidates the cache when data changes (e.g., after creating a product, "
     "the product list cache is evicted). Redis is configured with 10-minute TTL."),
    
    ("Q35: How did you handle concurrent stock updates?",
     "Concurrent stock updates are handled using pessimistic write locking. When a stock-in or stock-out operation "
     "is performed, the product row is locked with SELECT ... FOR UPDATE. This prevents two users from simultaneously "
     "reading the same stock quantity and making conflicting updates. If the lock can't be acquired, the request waits."),

    # Testing
    ("Q36: How did you write tests for this project?",
     "Unit tests use JUnit 5 and Mockito. Service tests mock repository dependencies and verify business logic. "
     "Controller integration tests use @WebMvcTest and mock services. Test profile uses H2 in-memory database. "
     "The pom.xml includes spring-boot-starter-test and spring-security-test dependencies."),

    # Specific Features
    ("Q37: How does the Dashboard aggregate data?",
     "DashboardService has a single @Cacheable method that: counts all products (productRepository.count()), "
     "sums total stock quantity (totalStockQuantity()), counts low-stock items (findLowStockProducts()), "
     "computes category distribution (countByCategory()), calculates monthly transaction trends (findTransactionsSince()), "
     "and fetches recent transactions (findTop10ByOrderByTransactionDateDesc()). All results are aggregated into a DashboardResponse."),
    
    ("Q38: How does the low-stock alert scheduling work?",
     "LowStockAlertService is annotated with @Scheduled(cron = '0 0 0 * * *') which runs the checkLowStock() "
     "method every day at midnight. It queries all products where quantity <= minimumStockLevel, logs the results, "
     "and could be extended to send email notifications. The @EnableScheduling annotation on the main application class "
     "enables scheduled task execution."),
    
    ("Q39: How is the audit trail implemented?",
     "AuditAspect is an @Aspect class that intercepts method calls on ProductService and InventoryService. "
     "It uses @Before advice on methods like createProduct, updateProduct, deleteProduct, stockIn, and stockOut. "
     "The aspect reads the current username from SecurityContextHolder and calls AuditLogService.log(username, action). "
     "This provides an automatic audit trail without modifying business logic."),

    # Database
    ("Q40: How is the database seeded with initial data?",
     "DataSeeder implements CommandLineRunner and runs after application startup. It checks if users already exist "
     "(userRepository.count() > 0). If not, it creates admin@inventra.com (ADMIN) and staff@inventra.com (STAFF), "
     "10 products across 3 categories, and 30 inventory transactions across the last 30 days. "
     "It's active for 'dev' and 'railway' profiles."),

    # Error Handling & Validation
    ("Q41: How did you handle validation errors?",
     "Two layers: (1) Backend uses Jakarta Bean Validation annotations (@NotBlank, @Email, @Size, @Min, @DecimalMin). "
     "When validation fails, MethodArgumentNotValidException is thrown. (2) GlobalExceptionHandler catches it and "
     "returns a 400 response with field-level error messages in a map. This gives clear feedback to the frontend."),

    # Frontend-Backend Communication
    ("Q42: How does the frontend know if the user is logged in after page refresh?",
     "The token is stored in localStorage (by default) or sessionStorage. When the app initializes, AuthService "
     "checks localStorage for the token. The auth guard calls isLoggedIn() which decodes the JWT payload, "
     "reads the exp claim, and compares it with the current time. If valid, the user stays logged in."),

    # Deployment
    ("Q43: How did you handle environment-specific configuration?",
     "Spring Profiles are used: 'dev' for local development (MySQL, DEBUG logging), 'test' for automated tests "
     "(H2 in-memory, no caching), 'railway' for Railway deployment (PostgreSQL, INFO logging, simple cache), "
     "and 'prod' for production (validate schema, no SHOW SQL). The active profile is set via SPRING_PROFILES_ACTIVE env var."),

    # More Technical
    ("Q44: What is the purpose of @Builder.Default on createdAt?",
     "@Builder.Default on createdAt = LocalDateTime.now() ensures that when an entity is created using the builder "
     "without explicitly setting createdAt, it defaults to the current time. Without this annotation, Lombok's builder "
     "would set createdAt to null if not explicitly provided."),
    
    ("Q45: How does the @PreUpdate annotation work?",
     "@PreUpdate is a JPA lifecycle callback that runs before an UPDATE operation. In the Product entity, "
     "the onUpdate() method sets updatedAt = LocalDateTime.now() to automatically track when a product was last modified. "
     "The user doesn't need to manually set this field."),

    # REST & API
    ("Q46: Why did you use ApiResponse wrapper for all responses?",
     "The generic ApiResponse<T> wrapper provides a consistent response format across all endpoints. It includes "
     "status code, message, data, and timestamp. This makes frontend error handling uniform – every response "
     "can be parsed the same way. The frontend's ApiResponse<T> interface mirrors this structure."),

    # Frontend specific
    ("Q47: How does the skeleton loading component work?",
     "The SkeletonComponent is a reusable component that displays placeholder animations while data is loading. "
     "It accepts a type input ('text', 'card', 'table-row', 'stat-card') to render different skeleton shapes. "
     "The parent component shows the skeleton when loading=true and hides it when data arrives."),

    # Design Decisions
    ("Q48: Why did you choose JWT over session-based authentication?",
     "JWT was chosen because: (1) It's stateless – no server-side session storage needed, making horizontal scaling easier. "
     "(2) It works well with REST APIs and mobile clients. (3) The token carries user information (email, role), "
     "reducing database lookups. (4) Expiry and validation are handled cryptographically."),
    
    ("Q49: Why is there a separate User entity that implements UserDetails?",
     "Implementing UserDetails directly on the User entity avoids an extra UserDetails wrapper class. "
     "Spring Security's DaoAuthenticationProvider works with any UserDetails implementation, so the User entity "
     "can be used directly. This is a clean approach for simple applications where the user entity maps directly "
     "to the security principal."),
    
    ("Q50: What would you improve if you had more time?",
     "Improvements: (1) Add real-time notifications via WebSocket for low-stock alerts. "
     "(2) Implement OAuth2 social login (Google/GitHub). (3) Add file upload for product images. "
     "(4) Generate PDF/CSV export for inventory reports. (5) Implement email notifications for low-stock. "
     "(6) Add pagination caching with Redis. (7) Improve test coverage with integration tests. "
     "(8) Add CI/CD with GitHub Actions for automated testing."),
]

for i, (q, a) in enumerate(questions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{q}')
    run.bold = True
    doc.add_paragraph(a)
    if i < len(questions):
        doc.add_paragraph()  # spacing

doc.add_page_break()

# ========================================================================
# SECTION 6: COMMON DESIGN DECISIONS
# ========================================================================
doc.add_heading('6. Common Design Decisions', level=1)

decisions = [
    ('Why Spring Boot 3.3.5?',
     'Spring Boot 3.x requires Java 17+ and supports Spring Framework 6. It has improved security, "
     "better GraalVM native image support, and modern Hibernate 6.x. Version 3.3.5 is a stable release with long-term support."),
    
    ('Why Angular 21?',
     'Angular 21 (the latest major) provides standalone components by default (no NgModules needed), "
     "improved server-side rendering, better Vite-based build system, and modern features like "
     "signals for reactive state management. It also has excellent TypeScript support."),
    
    ('Why JWT with 24-hour expiry?',
     "24 hours balances security with user experience. Shorter expiry (e.g., 15 minutes) would force users to "
     "re-login too frequently. Longer expiry increases the risk if a token is stolen. "
     "For a demo/educational project, 24 hours is reasonable."),
    
    ('Why BCrypt for password hashing?',
     "BCrypt is the industry standard for password hashing. It's intentionally slow (configurable cost factor) "
     "to resist brute-force and rainbow table attacks. Unlike MD5 or SHA-256, BCrypt includes a built-in salt "
     "to prevent identical passwords from producing the same hash."),
    
    ('Why PostgreSQL over MySQL?',
     "PostgreSQL was chosen because Railway provides native PostgreSQL support with automatic environment variable injection. "
     "PostgreSQL also offers advanced features like JSONB, CTEs, better concurrency control, and more robust query optimization. "
     "For the complex reporting queries in the dashboard and AI modules, PostgreSQL performs better."),
    
    ('Why Pessimistic Locking for stock transactions?',
     "Stock quantity is a critical business value. Optimistic locking (version column) would cause retry failures under high "
     "contention. Pessimistic locking with SELECT FOR UPDATE ensures that only one transaction can modify a product's stock "
     "at a time, preventing overselling or double-counting."),
    
    ('Why AOP for audit logging instead of manual logging?',
     "AOP (Aspect-Oriented Programming) keeps auditing concerns separate from business logic. Without AOP, every service "
     "method would need explicit audit log calls, cluttering the code. The AuditAspect intercepts methods declaratively, "
     "making the code cleaner and ensuring no audit log is accidentally missed."),
    
    ('Why Spring Cache abstraction with Redis/Simple fallback?',
     "Spring's @Cacheable annotation provides a declarative caching approach. The cache implementation is configurable: "
     "in production, Redis provides distributed caching; in simple deployments (like Railway without Redis), "
     "the 'simple' cache uses ConcurrentHashMap. This flexibility allows the code to be cache-implementation agnostic."),
    
    ('Why native queries for product search?',
     "The initial JPQL query used LOWER() for case-insensitive search, but PostgreSQL had columns stored as bytea type "
     "(a schema migration artifact). Native queries with ::text casts were needed to cast columns before applying LOWER(). "
     "Native queries also allow PostgreSQL-specific optimizations like ILIKE for case-insensitive matching."),
    
    ('Why first user gets ADMIN, rest get STAFF?',
     "This design simplifies onboarding. The first registered user (typically the business owner) gets ADMIN with "
     "full access. All subsequent users are STAFF with limited access (can view products and perform transactions "
     "but cannot create/edit/delete products or perform stock transactions). This eliminates the need for a separate "
     "admin creation flow."),
]

for title, desc in decisions:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    doc.add_paragraph(desc)

doc.add_page_break()

# ========================================================================
# SECTION 7: REDIS INTERVIEW QUESTIONS
# ========================================================================
doc.add_heading('7. Redis Interview Questions', level=1)

redis_qa = [
    ("Q1: What is Redis and why did you use it in this project?",
     "Redis (Remote Dictionary Server) is an in-memory data structure store used as a cache. We use Redis to cache "
     "dashboard aggregates, product lists, and low-stock queries. This reduces database load and improves API response times."),
    
    ("Q2: How is Redis configured in this project?",
     "RedisConfig.java is a @Configuration class with @ConditionalOnProperty(name='spring.cache.type', havingValue='redis'). "
     "It creates a RedisCacheManager with 10-minute TTL, no caching of null values, and JSON serialization "
     "using Jackson with JavaTimeModule. A RedisTemplate<String, Object> bean is also defined for general Redis operations."),
    
    ("Q3: What happens if Redis is not available?",
     "The project has a fallback: spring.cache.type=simple in the Railway profile. This uses a ConcurrentHashMap-based "
     "cache instead of Redis. The @Cacheable annotations work the same way regardless of the cache implementation. "
     "This allows the app to run without Redis during development or on platforms that don't support it."),
    
    ("Q4: What is the cache eviction strategy?",
     "Cache eviction is manual and transactional. When data changes (e.g., a product is created, updated, or deleted), "
     "@CacheEvict annotations on the service method clear the affected caches. For example, createProduct evicts "
     "'productList', 'dashboard', and 'lowStockProducts' caches. TTL-based expiration is also configured at 10 minutes."),
    
    ("Q5: How does @Cacheable work internally?",
     "Spring creates a proxy around the @Cacheable method. Before executing the method, it computes the cache key "
     "(using the SpEL expression in the `key` attribute) and checks if a value exists in the cache. If found, the "
     "cached value is returned without executing the method. If not found, the method executes, and the result is "
     "stored in the cache with the computed key."),
    
    ("Q6: What is the cache key for product listing?",
     "The key is a concatenation of all method parameters: "
     "'#page + '-' + #size + '-' + #sortBy + '-' + #sortDir + '-' + (#name ?: '') + '-' + (#category ?: '')'. "
     "This ensures different search/filter combinations have different cache entries."),
    
    ("Q7: How did you handle cache serialization?",
     "RedisConfig configures a Jackson2JsonRedisSerializer with ObjectMapper that has JavaTimeModule registered "
     "for LocalDateTime serialization. It also enables default typing (NonFinal) to include class information "
     "in the JSON. This ensures complex objects like DashboardResponse can be serialized/deserialized correctly."),
    
    ("Q8: What caching pattern is used — cache-aside or read-through?",
     "Spring's @Cacheable implements a read-through pattern. The application code doesn't manually check the cache. "
     "Instead, Spring intercepts the method call, checks the cache, executes the method if needed, and caches the result. "
     "This is cleaner than cache-aside pattern which requires explicit cache check logic in every method."),
    
    ("Q9: How would you handle cache stampede?",
     "Cache stampede happens when many requests try to rebuild the cache simultaneously after expiry. "
     "Solutions include: (1) Using @Cacheable with sync=true for basic locking. (2) Implementing early expiration "
     "with probabilistic refresh. (3) Using Redis SETNX for distributed locking. In this project, the 10-minute "
     "TTL and relatively low traffic make cache stampede unlikely."),
    
    ("Q10: What Redis data structures did you use?",
     "Primarily String-based caching via @Cacheable annotations. The RedisConfig also defines a RedisTemplate<String, Object> "
     "for future use cases like distributed locks (SETNX), rate limiting (INCR + EXPIRE), or session storage. "
     "Currently, all caching uses Spring's Cache Abstraction with String keys and JSON-serialized values."),
]

for q, a in redis_qa:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

doc.add_page_break()

# ========================================================================
# SECTION 8: JWT INTERVIEW QUESTIONS
# ========================================================================
doc.add_heading('8. JWT Interview Questions', level=1)

jwt_qa = [
    ("Q1: What is JWT and what are its components?",
     "JWT (JSON Web Token) is a compact, URL-safe token format used for authentication. It has three parts separated by dots: "
     "(1) Header – contains algorithm (HS256) and token type (JWT). (2) Payload – contains claims like subject (email), "
     "role, issued-at (iat), and expiration (exp). (3) Signature – verifies the token hasn't been tampered with."),
    
    ("Q2: How is the JWT signed and verified in this project?",
     "The JWT is signed using HMAC-SHA256 (HS256) algorithm. The secret key is derived from a Base64-encoded string "
     "configured in jwt.secret property. JJWT library (0.12.5) is used. JwtService.generateToken() creates the signature. "
     "JwtService.parseClaims() verifies the signature using the same secret key."),
    
    ("Q3: What claims are stored in the JWT payload?",
     "subject – user's email (used to load UserDetails), role – ADMIN or STAFF (for authorization), "
     "issuedAt – token creation timestamp, expiration – 24 hours from creation."),
    
    ("Q4: How does JwtAuthenticationFilter extract and validate the token?",
     "The filter extends OncePerRequestFilter. It reads the Authorization header, extracts the Bearer token, "
     "calls jwtService.extractEmail() to get the email, loads UserDetails from database, calls jwtService.isTokenValid(), "
     "and creates a UsernamePasswordAuthenticationToken which is set in SecurityContextHolder."),
    
    ("Q5: What happens if the JWT is expired?",
     "JwtService.isTokenValid() calls parseClaims() which internally validates the expiration. If expired, "
     "Jwts.parser() throws an ExpiredJwtException, which is caught by the catch block, and the method returns false. "
     "The filter then proceeds without setting authentication, causing the AuthorizationFilter to reject the request."),
    
    ("Q6: How did you handle JWT secret key rotation?",
     "Currently, the secret key is static and configured via the JWT_SECRET environment variable. For production, "
     "key rotation would require: (1) Supporting multiple keys with a key ID (kid) in the JWT header. "
     "(2) Storing keys in a secure vault. (3) Graceful overlap period where old keys are accepted but new tokens "
     "are issued with the new key. This is a future improvement."),
    
    ("Q7: What is the difference between JWT and OAuth2?",
     "JWT is a token format. OAuth2 is an authorization framework. JWT can be used within OAuth2 as the access token format. "
     "In this project, we use JWT directly for authentication (not OAuth2). OAuth2 would be needed for social login "
     "(Google, GitHub) or for allowing third-party apps to access the API."),
    
    ("Q8: How did you secure the JWT secret key?",
     "The secret key is configured via JWT_SECRET environment variable with a default fallback in application.yml. "
     "In production (Railway), the JWT_SECRET should be set in Railway's environment variables. "
     "The secret is Base64-encoded and must be at least 256 bits (32 bytes) for HS256 algorithm."),
    
    ("Q9: Can JWT tokens be revoked?",
     "JWT tokens are stateless and cannot be revoked server-side without additional infrastructure (like a token blacklist in Redis). "
     "For this reason, JWT expiration should be short. If revocation is needed, you could: (1) Maintain a Redis blacklist of revoked tokens. "
     "(2) Use short-lived tokens (15 minutes) with refresh tokens. (3) Implement token versioning in the database."),
    
    ("Q10: How did you prevent JWT token theft?",
     "Token theft mitigation: (1) Tokens are stored in localStorage (not cookies) to prevent CSRF. "
     "(2) HTTPS ensures tokens are encrypted in transit. (3) 24-hour expiry limits the damage window. "
     "(4) For higher security, tokens could be bound to the client's IP address or user-agent (token binding)."),
    
    ("Q11: What is the difference between Authentication and Authorization?",
     "Authentication verifies WHO the user is (identify verification via JWT token). Authorization determines WHAT "
     "the user can access (role-based access control). In this project, authentication is done by the JwtAuthenticationFilter, "
     "and authorization is done by the SecurityFilterChain rules and @PreAuthorize annotations."),
    
    ("Q12: How does the frontend remember the user across browser sessions?",
     "The token is stored in localStorage with 'rememberMe' option (default true). On app initialization, "
     "AuthService reads the token from localStorage. The auth guard calls isLoggedIn() which decodes the JWT "
     "and checks if it's still valid (not expired). If valid, the user is authenticated without re-login."),
]

for q, a in jwt_qa:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

doc.add_page_break()

# ========================================================================
# SECTION 9: AI INTEGRATION INTERVIEW QUESTIONS
# ========================================================================
doc.add_heading('9. AI Integration Interview Questions', level=1)

ai_qa = [
    ("Q1: How does the AI Assistant work in Inventra?",
     "The user types a natural language question in the chat interface. The frontend sends it to POST /api/ai/query. "
     "The AiService processes the question through GeminiService.classifyIntent() which sends it to Google Gemini API "
     "with a system prompt containing inventory context. Gemini returns the classified intent type. "
     "Based on the intent, AiService executes specific database queries and returns a response with summary + data."),
    
    ("Q2: What intents does the AI support?",
     "Supported intents: summary (overview of all products), low-stock (items below minimum stock level), "
     "stock-value (total inventory value), category-breakdown (products by category), product-search (search by name), "
     "product-count (total product count), and unknown (fallback for unrecognized queries)."),
    
    ("Q3: How did you integrate the Google Gemini API?",
     "GeminiService sends HTTP POST requests to the Gemini API endpoint (generativelanguage.googleapis.com/v1beta/models) "
     "using Spring's RestTemplate. The request includes the API key (configured via GEMINI_API_KEY env var), "
     "the model (gemini-2.0-flash), and a prompt that includes inventory context and the user's question."),
    
    ("Q4: What happens if the Gemini API key is not configured?",
     "If the API key is missing or invalid, the GeminiService catches the exception and falls back to a local "
     "classifyIntentLocally() method. This method uses simple keyword matching to determine the intent "
     "(e.g., 'low stock' -> low-stock intent, 'total value' -> stock-value intent)."),
    
    ("Q5: How did you handle API rate limiting?",
     "Currently, no rate limiting is implemented on the AI endpoint. For production, you would add: "
     "(1) Rate limiting per user (e.g., 10 requests/minute) using Redis INCR + EXPIRE. "
     "(2) Circuit breaker pattern using Resilience4j to stop calling Gemini if it's failing. "
     "(3) Request queuing for concurrent requests."),
    
    ("Q6: How did you prompt engineer the Gemini API?",
     "The system prompt includes: the current date, available product categories, product count, "
     "example products, and available intents with descriptions. This context helps Gemini accurately "
     "classify the user's intent. The prompt instructs Gemini to return a JSON response with intent, "
     "summary, and any relevant extracted parameters."),
    
    ("Q7: How did you handle the AI response format?",
     "Gemini returns a structured JSON object with fields: intent (the classified intent type), "
     "summary (a human-readable explanation), and parameters (extracted from the question). "
     "The AiService parses this JSON and executes the appropriate database query. "
     "The final response to the frontend includes the original question, intent, summary, and data."),
    
    ("Q8: How would you extend the AI to support more complex queries?",
     "To extend AI capabilities: (1) Add more intent types in the Gemini prompt. (2) Implement parameter extraction "
     "for filters (e.g., 'products in Electronics category under $100'). (3) Support multi-step reasoning "
     "(e.g., 'What is the total value of low-stock items?'). (4) Add vector embeddings for semantic search."),
    
    ("Q9: What are the costs and limitations of using Gemini API?",
     "Gemini 2.0 Flash is a cost-effective model with free tier limits. Costs include: per-character input/output pricing. "
     "Limitations: rate limits, latency (1-3 seconds per query), and model knowledge cutoff. "
     "To reduce costs, we cache common queries and use the local fallback for simple keyword-based classification."),
    
    ("Q10: How did you ensure the AI doesn't expose sensitive data?",
     "The AI is context-limited — it only has access to inventory data through the REST API. It cannot modify data. "
     "The Gemini API receives only the user's question and basic context (product categories, count). "
     "Database query execution is handled server-side by AiService, not by Gemini. "
     "This ensures no SQL injection or data exposure through the AI interface."),
]

for q, a in ai_qa:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

doc.add_page_break()

# ========================================================================
# SECTION 10: SPRING BOOT INTERVIEW QUESTIONS
# ========================================================================
doc.add_heading('10. Spring Boot Interview Questions', level=1)

spring_qa = [
    ("Q1: What is Spring Boot and why did you use it?",
     "Spring Boot is a framework that simplifies Spring application development by providing auto-configuration, "
     "embedded servers, and production-ready features. We used it because it reduces boilerplate configuration, "
     "provides seamless integration with Spring ecosystem (Security, Data JPA, Cache, AOP), and its embedded Tomcat "
     "makes deployment simple."),
    
    ("Q2: What is the difference between @Controller and @RestController?",
     "@Controller is used for MVC applications where views (templates) are returned. @RestController combines "
     "@Controller and @ResponseBody — it returns the response body directly (JSON/XML) without view resolution. "
     "In this project, all controllers use @RestController because they return JSON responses."),
    
    ("Q3: How does Spring Security's filter chain work?",
     "The SecurityFilterChain is a chain of servlet filters. Each filter performs a specific security function: "
     "CorsFilter handles CORS, CsrfFilter handles CSRF (disabled), JwtAuthFilter validates JWT tokens, "
     "AnonymousAuthenticationFilter creates anonymous auth, ExceptionTranslationFilter handles auth/access errors, "
     "and FilterSecurityInterceptor makes the final authorization decision based on the configured rules."),
    
    ("Q4: Explain the @Cacheable annotation workflow.",
     "@Cacheable is a Spring annotation that caches method return values. When a method with @Cacheable is called: "
     "(1) Spring generates a cache key using the SpEL expression in the `key` attribute. "
     "(2) It checks if a value exists in the specified cache. (3) If found, the cached value is returned. "
     "(4) If not found, the method executes, and the result is stored in the cache."),
    
    ("Q5: How does Spring Data JPA simplify database access?",
     "Spring Data JPA provides: (1) Automatic repository implementation — just extend JpaRepository and get CRUD methods. "
     "(2) Derived queries — method names like findByEmail() automatically generate SQL. "
     "(3) @Query for custom JPQL/native queries. (4) Pagination and sorting with Pageable parameter. "
     "(5) @Modifying for UPDATE/DELETE queries."),
    
    ("Q6: What is the purpose of @EnableScheduling?",
     "@EnableScheduling enables Spring's scheduled task execution capability. It's placed on the main application class. "
     "It allows the use of @Scheduled(cron = '...') annotations on methods. In this project, LowStockAlertService "
     "uses @Scheduled(cron = '0 0 0 * * *') to run the low-stock check every day at midnight."),
    
    ("Q7: Explain the @Transactional propagation and isolation levels.",
     "@Transactional manages database transactions. Propagation defines how transactions relate to each other "
     "(REQUIRED = use existing or create new, REQUIRES_NEW = always create new, SUPPORTS = optional). "
     "Isolation defines how transaction changes are visible to other transactions "
     "(READ_COMMITTED = default, prevents dirty reads; REPEATABLE_READ, SERIALIZABLE = stricter)."),
    
    ("Q8: What is the difference between @RequestParam and @PathVariable?",
     "@RequestParam extracts query parameters from the URL (e.g., /api/products?page=0&size=10). "
     "@PathVariable extracts values from the URL path template (e.g., /api/products/{id} captures the id from the URL). "
     "In this project, product listing uses @RequestParam for pagination/filtering, while product get/update/delete "
     "use @PathVariable for the product ID."),
    
    ("Q9: How does Spring Boot auto-configuration work?",
     "Spring Boot auto-configuration uses @Conditional annotations to automatically configure beans based on "
     "classpath dependencies, property values, and existing beans. For example, if spring-boot-starter-web is on the "
     "classpath, Spring Boot auto-configures Tomcat and DispatcherServlet. If H2 is on the classpath and no DataSource "
     "is configured, Spring Boot auto-configures an H2 in-memory DataSource."),
    
    ("Q10: What is AOP and how is it used in this project?",
     "AOP (Aspect-Oriented Programming) separates cross-cutting concerns from business logic. In this project, "
     "AuditAspect uses @Before advice to log user actions before product/inventory service methods execute. "
     "This keeps auditing logic separate from business logic, making both cleaner and easier to maintain."),
    
    ("Q11: How did you handle method-level security with @PreAuthorize?",
     "@EnableMethodSecurity enables method-level security. @PreAuthorize(\"hasRole('ADMIN')\") on controller methods "
     "ensures only users with ADMIN role can access those endpoints. The hasRole expression checks if the authenticated "
     "user's granted authorities include 'ROLE_ADMIN'. This provides defense-in-depth beyond the URL-based security rules."),
    
    ("Q12: What is the difference between @ComponentScan and @EnableAutoConfiguration?",
     "@ComponentScan tells Spring where to look for @Component, @Service, @Repository, @Controller beans. "
     "@EnableAutoConfiguration tells Spring Boot to automatically configure beans based on classpath dependencies. "
     "@SpringBootApplication includes both, plus @Configuration. Auto-configuration is what makes Spring Boot magic happen."),
    
    ("Q13: How does the CORS configuration work in this project?",
     "CORSConfigurationSource bean defines allowed origins (from app.cors.allowed-origins property), "
     "allowed methods (GET, POST, PUT, DELETE, OPTIONS), allowed headers (*), and allowCredentials(true). "
     "This is registered in the SecurityFilterChain via .cors(). In development, it allows localhost:4200. "
     "In production (Railway), it uses setAllowedOriginPatterns to support wildcards like https://*.railway.app."),
    
    ("Q14: What is the DispatcherServlet and its role?",
     "DispatcherServlet is the front controller in Spring MVC. It receives all HTTP requests and routes them to "
     "the appropriate handler (controller method). It uses handler mappings to find the right controller, "
     "handler adapters to invoke the method, and view resolvers to render responses. For REST controllers, "
     "it directly writes the response body using HttpMessageConverters."),
    
    ("Q15: How did you integrate multiple Spring Data modules (JPA + Redis)?",
     "Having both spring-boot-starter-data-jpa and spring-boot-starter-data-redis causes Spring to enter "
     "'strict repository configuration mode' as shown in the logs. Each repository interface is scanned for both "
     "JPA and Redis modules. To avoid conflicts, JPA repositories extend JpaRepository and Redis repositories "
     "would extend RedisRepository (none in this project). The Redis integration is only used via CacheManager."),
    
    ("Q16: What is the difference between PUT and PATCH in REST?",
     "PUT replaces the entire resource. In this project, updateProduct uses PUT with all product fields required. "
     "PATCH applies partial updates — only the fields that need to change are sent in the request body. "
     "For simplicity, this project uses PUT for all updates. PATCH would be better for real-world APIs "
     "to reduce data transfer and make partial updates explicit."),
    
    ("Q17: How would you add file upload support for product images?",
     "To add file upload: (1) Add @RequestParam MultipartFile on the controller method. "
     "(2) Store the file on disk or cloud storage (S3). (3) Save the file URL in the Product entity. "
     "(4) Configure max file size in application.yml (spring.servlet.multipart.max-file-size). "
     "Spring Boot's MultipartAutoConfiguration handles the file parsing automatically."),
    
    ("Q18: How does Spring handle JSON serialization?",
     "Spring uses Jackson's ObjectMapper (configured automatically) to serialize Java objects to JSON. "
     "@JsonIgnore can exclude fields. @JsonProperty can rename fields. The JavaTimeModule is needed for "
     "LocalDateTime serialization. In this project, the @Data annotation (Lombok) generates getters which Jackson "
     "uses to determine which properties to serialize."),
    
    ("Q19: What is the role of application.properties/yml?",
     "application.yml is the main configuration file for Spring Boot. It configures: server port, datasource, "
     "JPA settings, logging levels, external service API keys, JWT secret/expiration, and profile-specific overrides. "
     "Environment variables can override these values at runtime (${VAR_NAME:default} syntax), "
     "which is essential for production deployment without hard-coding secrets."),
    
    ("Q20: How did you test the Spring Boot application?",
     "Testing uses: (1) JUnit 5 — test framework. (2) Mockito — mocking dependencies in unit tests. "
     "(3) @SpringBootTest — integration tests with full application context. "
     "(4) @WebMvcTest — controller layer tests with mocked services. "
     "(5) H2 in-memory database for repository tests. "
     "(6) Test profile (application-test.yml) with H2 and no caching."),
]

for q, a in spring_qa:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

# ---------- FOOTER ----------
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generated on {datetime.date.today().strftime("%B %d, %Y")} for Chetana Mahajan')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ---------- SAVE ----------
output_path = r'D:\Symbosis Traning\1 Workspace\1 Projects\Inventra – AI-Powered Inventory Management System\Inventra_Complete_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total sections: 10')
print(f'Total interview Q&A: {len(questions) + len(redis_qa) + len(jwt_qa) + len(ai_qa) + len(spring_qa)}')
